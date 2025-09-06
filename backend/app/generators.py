from typing import List, Dict
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from io import BytesIO

def make_docx(resume_text: str, job_title: str|None, suggestions: Dict[str, List[str]]) -> bytes:
    doc = Document()
    h = doc.add_heading("Tailored Resume (Draft)", level=1)
    if job_title:
        doc.add_paragraph(f"Target Role: {job_title}")
    doc.add_paragraph(" ")
    doc.add_heading("Suggestions to Weave Into Bullets", level=2)
    doc.add_paragraph("Add:")
    for s in suggestions.get("add", []):
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph("Emphasize:")
    for s in suggestions.get("emphasize", []):
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph(" ")
    doc.add_heading("Original Resume Text (for editing)", level=2)
    p = doc.add_paragraph(resume_text)
    p.style.font.size = Pt(10)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def make_pdf(resume_text: str, job_title: str|None, suggestions: Dict[str, List[str]]) -> bytes:
    buff = BytesIO()
    c = canvas.Canvas(buff, pagesize=LETTER)
    width, height = LETTER
    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Tailored Resume (Draft)")
    y -= 20
    c.setFont("Helvetica", 11)
    if job_title:
        c.drawString(50, y, f"Target Role: {job_title}")
        y -= 20
    c.setFont("Helvetica-Bold", 12); c.drawString(50, y, "Suggestions - Add"); y -= 16
    c.setFont("Helvetica", 10)
    for s in suggestions.get("add", []):
        for line in wrap(s, 90):
            c.drawString(60, y, f"• {line}"); y -= 12
            if y < 70: c.showPage(); y = height - 50
    c.setFont("Helvetica-Bold", 12); c.drawString(50, y, "Suggestions - Emphasize"); y -= 16
    c.setFont("Helvetica", 10)
    for s in suggestions.get("emphasize", []):
        for line in wrap(s, 90):
            c.drawString(60, y, f"• {line}"); y -= 12
            if y < 70: c.showPage(); y = height - 50
    c.setFont("Helvetica-Bold", 12); c.drawString(50, y, "Original Resume Text"); y -= 16
    c.setFont("Helvetica", 9)
    for line in wrap(resume_text, 100):
        c.drawString(50, y, line); y -= 11
        if y < 70: c.showPage(); y = height - 50
    c.save()
    return buff.getvalue()

def wrap(text: str, width: int):
    # tiny word wrapper
    words = text.split()
    line = []
    count = 0
    for w in words:
        if count + len(w) + (1 if line else 0) > width:
            yield " ".join(line); line = [w]; count = len(w)
        else:
            line.append(w); count += len(w) + (1 if line else 0)
    if line: yield " ".join(line)
